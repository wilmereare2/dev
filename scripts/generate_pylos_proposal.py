"""Generate Pylos Game client proposal DOCX."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date

OUTPUT = r"D:\op\Pylos_Game_Development_Proposal.docx"


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def main():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PYLOS MOBILE GAME\nDevelopment Proposal")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Prepared: {date.today().strftime('%B %d, %Y')}\nProject ID: 40686055")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "Thank you for sharing the Pylos game requirements and rule documentation. "
        "Based on your detailed specification document, this project is a full-featured "
        "mobile game platform — not a simple game rebuild — modeled after your existing "
        "chess application with online play, monetization, social features, and progression systems."
    )
    doc.add_paragraph(
        "This proposal outlines a phased development approach with realistic timelines and "
        "budget estimates, allowing you to launch in stages and scale features over time."
    )

    # 2. Project Understanding
    add_heading(doc, "2. Project Understanding", 1)
    doc.add_paragraph("We understand you need:")
    items = [
        ("Platform: ", "iOS and Android (Unity engine)"),
        ("Situation: ", "Original source code lost — full rebuild required"),
        ("Publishing: ", "App Store and Google Play Store deployment support"),
        ("Reference: ", "Feature parity with your existing chess game platform"),
        ("Game: ", "Pylos — strategic pyramid-building board game by David G. Royffe"),
    ]
    for bold, text in items:
        add_bullet(doc, text, bold)

    # 3. Full Scope
    add_heading(doc, "3. Complete Feature Scope (From Your Document)", 1)

    add_heading(doc, "3.1 Core Game Modes", 2)
    modes = [
        "Children Version — simplified rules (move-up only, no square reclaim)",
        "Mature Version — full official rules with square formation and alignment reclaim",
        "Maniac Version — expanded 7×7 board with 140 pieces (70 per player)",
        "Game timers: 10 minutes (4×4) / 20 minutes (7×7)",
    ]
    for m in modes:
        add_bullet(doc, m)

    add_heading(doc, "3.2 Gameplay & AI", 2)
    for item in [
        "Play vs Computer — Easy, Medium, Hard, Practice Match",
        "Play vs Friends — invite system with purchasable friend slots (diamonds)",
        "Play Online — ranked matchmaking for all three rule sets",
        "Interactive tutorial and in-app rulebook for each mode",
        "Power-ups / ammos — freeze opponent timer, board obstacles, dirt tiles (IAP + VIP perks)",
        "Tournaments — unlock after 20 Mature games; Maniac mode gated behind progression",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.3 Monetization & Economy", 2)
    for item in [
        "VIP Membership — 7-day free trial, then $5/week, $10/month, or $50/year",
        "VIP benefits — ad removal, exclusive emojis, VIP character, weekly diamonds, exclusive board skins",
        "In-app purchases — character skins, marble designs, board themes, power-ups",
        "Rewarded ads — watch video for gift boxes (as in chess app)",
        "Chest & key system — win chests on victory; keys unlock random rewards (League of Legends style)",
        "Thematic environments — desert, jungle, space, etc.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.4 Player Profile & Social", 2)
    for item in [
        "Google Play / Game Center account login",
        "Leaderboards, match history, social features",
        "Quests, achievements, and challenges",
        "Collection system — skins, pets, relics",
        "Progression and experience-based unlocks",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.5 Future Enhancements (Phase 2+)", 2)
    for item in [
        "Gravity mechanic — marbles restricted to specific levels",
        "Color-coded spheres with placement restrictions",
        "Team play and cooperative modes",
        "Battle royale multiplayer mode",
    ]:
        add_bullet(doc, item)

    # 4. Phased Approach
    add_heading(doc, "4. Recommended Phased Development", 1)
    doc.add_paragraph(
        "Given the full scope, we recommend building in four phases. Each phase delivers "
        "a working, testable product before moving to the next."
    )

    phases = [
        (
            "Phase 1 — Core Game MVP",
            "6–8 weeks",
            "₹3,50,000 – ₹4,50,000",
            [
                "Classic Pylos (Mature rules) — full game logic",
                "Local 2-player (pass-and-play)",
                "AI opponent — Easy, Medium, Hard",
                "Clean mobile UI with board, marbles, and animations",
                "Interactive tutorial and rulebook",
                "Basic menu and settings",
                "iOS + Android builds (TestFlight / internal testing)",
            ],
        ),
        (
            "Phase 2 — Online & Accounts",
            "6–8 weeks",
            "₹4,00,000 – ₹5,50,000",
            [
                "User accounts — Google Play / Apple sign-in",
                "Online multiplayer and matchmaking",
                "Play vs Friends with invite system",
                "Match history and basic player profiles",
                "Game timers (10 min / 20 min)",
                "Children Version game mode",
                "Backend server setup and deployment",
            ],
        ),
        (
            "Phase 3 — Monetization & Progression",
            "4–6 weeks",
            "₹3,00,000 – ₹4,00,000",
            [
                "Ad integration (banner, interstitial, rewarded video)",
                "VIP subscription — trial, weekly/monthly/yearly plans",
                "In-app purchases — skins, boards, characters",
                "Chest & key reward system",
                "Quests, achievements, and daily rewards",
                "Gift box via rewarded video ads",
                "Diamond economy and shop",
            ],
        ),
        (
            "Phase 4 — Advanced Features",
            "4–6 weeks",
            "₹3,50,000 – ₹5,00,000",
            [
                "Maniac Version (7×7 board, 140 pieces)",
                "Power-ups and special ammos (IAP)",
                "Tournaments and ranked seasons",
                "Leaderboards and social features",
                "Collection — pets, relics, themed environments",
                "Friend slots (diamond purchase)",
                "App Store + Play Store publishing support",
            ],
        ),
    ]

    for name, duration, price, deliverables in phases:
        add_heading(doc, name, 2)
        p = doc.add_paragraph()
        r1 = p.add_run(f"Duration: {duration}    |    Estimated Cost: {price}")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)
        doc.add_paragraph("Deliverables:")
        for d in deliverables:
            add_bullet(doc, d)

    # 5. Timeline Summary
    add_heading(doc, "5. Timeline Summary", 1)

    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Phase", "Duration", "Cumulative", "Estimated Cost (INR)"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1A375E")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows_data = [
        ("Phase 1 — Core MVP", "6–8 weeks", "6–8 weeks", "₹3,50,000 – ₹4,50,000"),
        ("Phase 2 — Online", "6–8 weeks", "12–16 weeks", "₹4,00,000 – ₹5,50,000"),
        ("Phase 3 — Monetization", "4–6 weeks", "16–22 weeks", "₹3,00,000 – ₹4,00,000"),
        ("Phase 4 — Advanced", "4–6 weeks", "20–28 weeks", "₹3,50,000 – ₹5,00,000"),
    ]
    for ri, row in enumerate(rows_data, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val

    doc.add_paragraph()
    total = doc.add_paragraph()
    r = total.add_run(
        "Full Project Total: 5–7 months  |  ₹14,00,000 – ₹19,00,000 INR"
    )
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

    # 6. Budget Note
    add_heading(doc, "6. Budget Alignment Note", 1)
    doc.add_paragraph(
        "Your posted project budget (₹1,000 – ₹21,000 INR) covers approximately Phase 1 "
        "planning and a basic prototype only. The complete feature set described in your "
        "document requires the phased budget outlined above."
    )
    doc.add_paragraph(
        "Recommended starting point: Phase 1 (Core MVP) — delivers a playable, "
        "publishable-quality game with AI that you can demo and use to secure further "
        "investment or proceed with later phases."
    )

    # 7. What's Included / Excluded
    add_heading(doc, "7. Included vs. Not Included", 1)

    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "Included in Development"
    table2.rows[0].cells[1].text = "Client Responsibility"
    for cell in table2.rows[0].cells:
        set_cell_shading(cell, "E8EEF4")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    included = [
        "Game design and development (Unity)",
        "All game logic per official Pylos rules",
        "UI/UX design and implementation",
        "AI opponent programming",
        "Backend server (Phases 2+)",
        "iOS and Android builds",
        "Store listing setup assistance",
        "Bug fixes during each phase",
        "Source code handover on completion",
    ]
    client_resp = [
        "Apple Developer account ($99/year)",
        "Google Play Developer account ($25 one-time)",
        "App Store / Play Store listing fees",
        "Server hosting costs (ongoing)",
        "Marketing and user acquisition",
        "Original art assets if custom branding required",
        "Legal review (IP, privacy policy, terms of service)",
    ]

    for inc, cli in zip(included, client_resp):
        row = table2.add_row()
        row.cells[0].text = inc
        row.cells[1].text = cli

    # 8. Requirements from Client
    add_heading(doc, "8. Information Needed From You", 1)
    for item in [
        "Screenshots or video of your previous Pylos app (if available)",
        "Access to your chess app for feature reference (APK or TestFlight link)",
        "Preferred visual style — classic wood look vs. modern theme",
        "Confirmation of which phases to proceed with",
        "Apple and Google developer account status",
        "Target launch date",
        "Payment milestone preferences (we suggest per-phase payments)",
    ]:
        add_bullet(doc, item)

    # 9. Payment Terms
    add_heading(doc, "9. Suggested Payment Structure", 1)
    for item in [
        "40% upfront at phase start",
        "30% on delivery of working milestone build",
        "30% on phase completion and source code handover",
        "Each phase is a separate agreement — you can stop after any phase",
    ]:
        add_bullet(doc, item)

    # 10. Next Steps
    add_heading(doc, "10. Next Steps", 1)
    steps = [
        "Review this proposal and confirm desired phase(s)",
        "Share reference materials (old app screenshots, chess app access)",
        "Agree on Phase 1 scope, timeline, and fixed price",
        "Sign agreement and initiate Phase 1 development",
        "Weekly progress updates via chat with playable builds every 2 weeks",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = closing.add_run(
        "We are ready to start Phase 1 immediately upon confirmation.\n"
        "Looking forward to building PYLOS with you."
    )
    r.italic = True
    r.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
